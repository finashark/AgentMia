"""
File Service
Handles file operations for scripts (.docx and .txt)
"""
import os
from docx import Document
from datetime import datetime
from typing import Optional
from config import SCRIPT_FOLDER, SUPPORTED_FILE_FORMATS

class FileService:
    def __init__(self):
        """Initialize file service"""
        self.script_folder = SCRIPT_FOLDER
        
        # Create Script Folder if it doesn't exist
        if not os.path.exists(self.script_folder):
            os.makedirs(self.script_folder)
    
    def read_file(self, file_path: str) -> str:
        """
        Read content from a file (.txt or .docx)
        
        Args:
            file_path: Path to the file
            
        Returns:
            File content as string
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in SUPPORTED_FILE_FORMATS:
            raise ValueError(f"Định dạng file không được hỗ trợ. Chỉ hỗ trợ: {', '.join(SUPPORTED_FILE_FORMATS)}")
        
        try:
            if file_ext == '.txt':
                return self._read_txt(file_path)
            elif file_ext == '.docx':
                return self._read_docx(file_path)
        except Exception as e:
            raise Exception(f"Lỗi khi đọc file: {str(e)}")
    
    def _read_txt(self, file_path: str) -> str:
        """Read .txt file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _read_docx(self, file_path: str) -> str:
        """Read .docx file"""
        doc = Document(file_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        return '\n\n'.join(content)
    
    def read_uploaded_file(self, uploaded_file) -> str:
        """
        Read content from Streamlit uploaded file
        
        Args:
            uploaded_file: Streamlit UploadedFile object
            
        Returns:
            File content as string
        """
        file_name = uploaded_file.name
        file_ext = os.path.splitext(file_name)[1].lower()
        
        if file_ext not in SUPPORTED_FILE_FORMATS:
            raise ValueError(f"Định dạng file không được hỗ trợ. Chỉ hỗ trợ: {', '.join(SUPPORTED_FILE_FORMATS)}")
        
        try:
            if file_ext == '.txt':
                return uploaded_file.read().decode('utf-8')
            elif file_ext == '.docx':
                doc = Document(uploaded_file)
                content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text)
                return '\n\n'.join(content)
        except Exception as e:
            raise Exception(f"Lỗi khi đọc file upload: {str(e)}")
    
    def save_script(self, content: str, filename: Optional[str] = None, format: str = 'txt') -> str:
        """
        Save script content to Script Folder
        
        Args:
            content: Script content to save
            filename: Optional filename (without extension)
            format: File format ('txt' or 'docx')
            
        Returns:
            Path to saved file
        """
        if format not in ['txt', 'docx']:
            raise ValueError("Format phải là 'txt' hoặc 'docx'")
        
        # Generate filename if not provided
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"script_{timestamp}"
        
        # Remove extension if provided in filename
        filename = os.path.splitext(filename)[0]
        
        # Create full path
        file_path = os.path.join(self.script_folder, f"{filename}.{format}")
        
        try:
            if format == 'txt':
                self._save_txt(file_path, content)
            elif format == 'docx':
                self._save_docx(file_path, content)
            
            return file_path
        except Exception as e:
            raise Exception(f"Lỗi khi lưu script: {str(e)}")
    
    def _save_txt(self, file_path: str, content: str):
        """Save content as .txt file"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    def _save_docx(self, file_path: str, content: str):
        """Save content as .docx file"""
        doc = Document()
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(file_path)
    
    def list_scripts(self) -> list:
        """
        List all saved scripts in Script Folder
        
        Returns:
            List of script filenames
        """
        try:
            files = []
            for file in os.listdir(self.script_folder):
                file_ext = os.path.splitext(file)[1].lower()
                if file_ext in SUPPORTED_FILE_FORMATS:
                    file_path = os.path.join(self.script_folder, file)
                    file_stat = os.stat(file_path)
                    files.append({
                        'name': file,
                        'path': file_path,
                        'size': file_stat.st_size,
                        'modified': datetime.fromtimestamp(file_stat.st_mtime)
                    })
            
            # Sort by modified date (newest first)
            files.sort(key=lambda x: x['modified'], reverse=True)
            return files
        except Exception as e:
            raise Exception(f"Lỗi khi liệt kê scripts: {str(e)}")
    
    def delete_script(self, file_path: str) -> bool:
        """
        Delete a script file
        
        Args:
            file_path: Path to the file to delete
            
        Returns:
            True if successful
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
            return False
        except Exception as e:
            raise Exception(f"Lỗi khi xóa script: {str(e)}")

# Test function
if __name__ == "__main__":
    service = FileService()
    
    # Test save
    test_content = "Đây là một script test\n\nNội dung dòng 2"
    saved_path = service.save_script(test_content, "test_script", "txt")
    print(f"Saved to: {saved_path}")
    
    # Test list
    scripts = service.list_scripts()
    print(f"Found {len(scripts)} scripts")
